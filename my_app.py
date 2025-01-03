from docx import Document
from docx.shared import Inches
import pyttsx3


def speak(text):
    pyttsx3.speak(text)


document = Document()

# name, phone number and email details
name = input("What is your name? ")
# speak(f"Hello {name}, how are you today? Please go on")
phone_number = input("What is your phone number? ")
# speak(f"Your phone number is {phone_number}")
email = input("What is your email? ")
# speak(f"The email you inserted is {email}")

# profile picture
document.add_picture(
    f"./pictures/{name}.jpg"
    width=Inches(2.0)
)

document.add_paragraph(
    name + " | " + phone_number + " | " + email
)

# about me
document.add_heading("About me")
about_me = input("Tell me about yourself: ")
document.add_paragraph(about_me)

# work experience
document.add_heading("Work Experience")
p = document.add_paragraph()

company = input("Enter company: ")
from_date = input("From Date: ")
to_date = input("To Date: ")

p.add_run(company + " ").bold = True
p.add_run(from_date + "-" + to_date + "\n").italic = True

experience_details = input(
    "Describe your experience at " + company + " "
)
p.add_run(experience_details)

# more experiences
while True:
    has_more_experiences = input(
        "Do you have more experiences? Yes(Y) or No(N) ")
    if has_more_experiences.lower() == "yes":
        p = document.add_paragraph()

        company = input("Enter company: ")
        from_date = input("From Date: ")
        to_date = input("To Date: ")

        p.add_run(company + " ").bold = True
        p.add_run(from_date + "-" + to_date + "\n").italic = True

        experience_details = input(
            "Describe your experience at " + company + " "
        )
        p.add_run(experience_details)
    else:
        break

# skills
document.add_heading("Skills")
skill = input("Insert the first skill: ")
s = document.add_paragraph(skill)
s.style = "List Bullet"

# insert more skills
while True:
    has_more_skills = input(
        "Do you have more skills to add? Yes(Y) or No(N) ")
    if has_more_skills.lower() == "yes":
        skill = input("Enter skill: ")
        s = document.add_paragraph(skill)
        s.style = "List Bullet"
    else:
        break

# footer
section = document.sections[0]
footer = section.footer
p = footer.paragraphs[0]
p.text = "CV generated with Python programming"

document.save(f"cv_{name}.docx")  # create the cv document
